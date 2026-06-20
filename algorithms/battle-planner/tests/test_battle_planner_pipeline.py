from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from battle_planner.io.enemy_reader import read_enemy_situation
from battle_planner.pipeline import PlanningPipeline


def _write_json(path: Path, payload: dict) -> Path:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _mock_config(path: Path) -> Path:
    return _write_json(
        path,
        {
            "llm": {"provider": "mock"},
            "algorithm": {
                "reserve_ratio": 0.1,
                "max_group_size": 4,
                "default_air_assault_personnel": 18,
            },
        },
    )


def _bare_threat(path: Path) -> Path:
    return _write_json(
        path,
        {
            "ok": True,
            "threatLevel": "高",
            "threatScore": 82,
            "targetAssessments": [
                {
                    "id": "target-001",
                    "name": "东侧防空节点",
                    "category": "air_defense",
                    "threatScore": 86,
                    "valueScore": 78,
                    "location": {"coordinates": [118.12, 32.04], "locationDescription": "东侧高地"},
                    "sourceTarget": {"subCategory": "manportable_air_defense"},
                },
                {
                    "id": "target-002",
                    "name": "北侧通信中继站",
                    "category": "command_control",
                    "threatScore": 64,
                    "valueScore": 88,
                    "location": {"coordinates": [118.22, 32.1], "locationDescription": "北侧山脊"},
                    "sourceTarget": {"subCategory": "communication_relay"},
                },
            ],
        },
    )


def _wrapped_threat(path: Path) -> Path:
    bare = json.loads(_bare_threat(path.with_name("bare-source.json")).read_text(encoding="utf-8"))
    return _write_json(
        path,
        {
            "schemaVersion": "planning-artifact-export-v1",
            "generatedAt": "2026-06-20T00:00:00.000Z",
            "step": {"algorithmId": "enemy-threat-analysis"},
            "artifact": {"name": "敌情威胁结构化结果"},
            "output": bare,
        },
    )


def _friendly_txt(path: Path) -> Path:
    path.write_text(
        "\n".join(
            [
                "二型武装直升机: 8 架",
                "运输直升机: 4 架",
                "侦察直升机: 2 架",
                "空地导弹: 32 枚",
                "火箭弹: 180 发",
                "航炮弹: 4200 发",
                "机降突击人员: 60 名",
            ]
        ),
        encoding="utf-8",
    )
    return path


def _friendly_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("二型武装直升机: 6 架")
    document.add_paragraph("运输直升机: 3 架")
    document.add_paragraph("空地导弹: 24 枚")
    document.add_paragraph("火箭弹: 120 发")
    document.add_paragraph("机降突击人员: 48 名")
    document.save(path)
    return path


def _friendly_without_weapons_json(path: Path) -> Path:
    return _write_json(
        path,
        {
            "friendly_forces": {
                "helicopters": [
                    {
                        "model": "二型武装直升机",
                        "role": "armed",
                        "available": 4,
                        "capabilities": ["防空压制", "火力打击", "火力压制", "护航"],
                        "weapon_capacity": {"空地导弹": 4, "火箭弹": 16},
                        "personnel_capacity": 0,
                    },
                    {
                        "model": "运输直升机",
                        "role": "transport",
                        "available": 3,
                        "capabilities": ["机降突击", "人员输送"],
                        "weapon_capacity": {},
                        "personnel_capacity": 12,
                    },
                ],
                "weapons": [],
                "personnel": [{"role": "机降突击人员", "available": 24}],
                "constraints": {"preserve_reserve": False},
            }
        },
    )


def test_mock_pipeline_runs_with_txt_friendly_document(tmp_path: Path) -> None:
    pipeline = PlanningPipeline.from_config(str(_mock_config(tmp_path / "config.json")))
    result = pipeline.run(
        enemy_files=[_wrapped_threat(tmp_path / "threat.json")],
        friendly_files=[_friendly_txt(tmp_path / "friendly.txt")],
        output_dir=tmp_path / "outputs",
    )

    assert result.total_groups >= 1
    assert result.task_groups
    assert (tmp_path / "outputs" / "grouping_result.json").exists()
    assert result.metadata["llm_provider"] == "mock"


def test_enemy_reader_accepts_planning_artifact_wrapper(tmp_path: Path) -> None:
    situation = read_enemy_situation(_wrapped_threat(tmp_path / "wrapped-threat.json"))

    assert len(situation.targets) == 2
    assert situation.metadata["schemaVersion"] == "planning-artifact-export-v1"
    assert situation.targets[0].name == "东侧防空节点"
    assert situation.targets[0].coordinate.lon == 118.12


def test_mock_pipeline_runs_with_docx_friendly_document(tmp_path: Path) -> None:
    pipeline = PlanningPipeline.from_config(str(_mock_config(tmp_path / "config.json")))
    result = pipeline.run(
        enemy_files=[_wrapped_threat(tmp_path / "threat.json")],
        friendly_files=[_friendly_docx(tmp_path / "friendly.docx")],
    )

    assert result.total_groups >= 1
    assert result.task_groups[0].platforms


def test_fire_strike_without_loaded_weapons_has_zero_firepower(tmp_path: Path) -> None:
    pipeline = PlanningPipeline.from_config(str(_mock_config(tmp_path / "config.json")))
    result = pipeline.run(
        enemy_files=[_wrapped_threat(tmp_path / "threat.json")],
        friendly_files=[_friendly_without_weapons_json(tmp_path / "friendly-no-weapons.json")],
    )

    fire_group = next(group for group in result.task_groups if group.task_type in {"防空压制", "火力打击", "火力压制", "通信压制", "破袭打击"})

    assert fire_group.firepower_score == 0
    assert fire_group.firepower_breakdown["combinedFirepower"] == 0
    assert fire_group.firepower_breakdown["weaponEquipmentPower"] == 0
    assert fire_group.firepower_breakdown["hasLoadedWeapon"] is False
    assert fire_group.strike_weapon_requirement_met is False
    assert fire_group.assignment_eligible_for_strike is False
    assert any(issue.severity == "error" and "未形成实际武器装载" in issue.message for issue in fire_group.issues)
