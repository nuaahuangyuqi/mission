"""Friendly-force document readers."""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Union

from docx import Document


@dataclass
class FriendlyDocument:
    path: str
    source_type: str
    content: str
    structured_hint: Optional[Dict[str, Any]] = field(default=None)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_json(path: Path) -> FriendlyDocument:
    with path.open("r", encoding="utf-8") as file:
        raw = json.load(file)
    content = json.dumps(raw, ensure_ascii=False, indent=2)
    return FriendlyDocument(path=str(path), source_type="json", content=content, structured_hint=raw)


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_friendly_document(path: Union[str, Path]) -> FriendlyDocument:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".json":
        return _read_json(source)
    if suffix in {".txt", ".md", ".markdown"}:
        return FriendlyDocument(path=str(source), source_type=suffix.lstrip("."), content=_read_text(source))
    if suffix == ".docx":
        return FriendlyDocument(path=str(source), source_type="docx", content=_read_docx(source))
    raise ValueError(f"不支持的己方信息文档格式: {source.suffix}")


def read_friendly_documents(paths: Sequence[Union[str, Path]]) -> FriendlyDocument:
    if not paths:
        raise ValueError("至少需要一个己方信息文档")
    documents = [read_friendly_document(path) for path in paths]
    if len(documents) == 1:
        return documents[0]

    content_parts: List[str] = []
    source_types: List[str] = []
    source_paths: List[str] = []
    structured_hints: List[Dict[str, Any]] = []
    for index, document in enumerate(documents, start=1):
        source_types.append(document.source_type)
        source_paths.append(document.path)
        content_parts.append(
            "\n".join(
                [
                    f"--- 己方信息文档 {index} ---",
                    f"来源：{document.path}",
                    f"类型：{document.source_type}",
                    document.content,
                ]
            )
        )
        if document.structured_hint:
            structured_hints.append(document.structured_hint)

    return FriendlyDocument(
        path=";".join(source_paths),
        source_type="+".join(source_types),
        content="\n\n".join(content_parts),
        structured_hint=_merge_structured_hints(structured_hints) if structured_hints else None,
    )


def _merge_structured_hints(hints: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    merged_forces: Dict[str, Any] = {
        "helicopters": [],
        "weapons": [],
        "personnel": [],
        "task_capabilities": [],
        "grouping_rules": {},
        "constraints": {},
        "warnings": [],
    }
    for hint in hints:
        forces = hint.get("friendly_forces", hint)
        if not isinstance(forces, dict):
            continue
        merged_forces["helicopters"].extend(forces.get("helicopters") or [])
        merged_forces["weapons"].extend(forces.get("weapons") or [])
        merged_forces["personnel"].extend(forces.get("personnel") or [])
        merged_forces["task_capabilities"].extend(forces.get("task_capabilities") or [])
        merged_forces["grouping_rules"].update(forces.get("grouping_rules") or {})
        merged_forces["constraints"].update(forces.get("constraints") or {})
        merged_forces["warnings"].extend(forces.get("warnings") or [])

    merged_forces["helicopters"] = _merge_by_name(merged_forces["helicopters"], "model", "available")
    merged_forces["weapons"] = _merge_by_name(merged_forces["weapons"], "name", "available")
    merged_forces["personnel"] = _merge_by_name(merged_forces["personnel"], "role", "available")
    merged_forces["task_capabilities"] = sorted({str(item) for item in merged_forces["task_capabilities"]})
    return {"friendly_forces": merged_forces}


def _merge_by_name(items: Sequence[Any], key_field: str, count_field: str) -> List[Dict[str, Any]]:
    merged: Dict[str, Dict[str, Any]] = {}
    for item in items:
        if not isinstance(item, dict):
            continue
        key = str(item.get(key_field, "")).strip()
        if not key:
            continue
        if key not in merged:
            merged[key] = dict(item)
            continue
        current = merged[key]
        current[count_field] = int(current.get(count_field, 0) or 0) + int(item.get(count_field, 0) or 0)
        for field in ("capabilities", "effects"):
            if field in item or field in current:
                current[field] = sorted({str(value) for value in current.get(field, []) + item.get(field, [])})
        if key_field == "model":
            current.setdefault("weapon_capacity", item.get("weapon_capacity", {}))
            current.setdefault("personnel_capacity", item.get("personnel_capacity", 0))
    return list(merged.values())
