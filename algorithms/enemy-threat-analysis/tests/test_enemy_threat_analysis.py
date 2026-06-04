from __future__ import annotations

import base64
import io
import json
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from enemy_threat_analysis import analyze
from enemy_threat_analysis.errors import AnalysisInputError
from enemy_threat_analysis.image_exporter import IMAGE_SIZE, build_heatmap_base64
from enemy_threat_analysis.llm_extractor import parse_llm_extraction


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_EXTRACTION = ROOT / "examples" / "mock_extraction.json"


REQUIRED_FIELDS = {
    "implementationStatus",
    "builtinMethodKey",
    "builtinMethodLabel",
    "appliedOptions",
    "inputSummary",
    "evidenceTrace",
    "threatLevel",
    "threatScore",
    "enemyUnitCount",
    "identifiedThreatNodeCount",
    "enemyIntentions",
    "deploymentSectors",
    "fireCoverage",
    "airDefenseSystem",
    "reconEarlyWarning",
    "antiAirborneFacilities",
    "impactAnalysis",
    "visualization",
}


def load_sample() -> dict:
    return json.loads(SAMPLE_EXTRACTION.read_text(encoding="utf-8"))


def test_empty_input_requires_files_when_no_extraction() -> None:
    with pytest.raises(AnalysisInputError):
        analyze([], generate_assessment=False)


def test_structured_output_contains_required_fields() -> None:
    result = analyze(
        None,
        analysis_focus="comprehensive",
        heatmap_density="low",
        impact_bias="balanced",
        extraction_json=load_sample(),
        generate_assessment=False,
    )
    assert REQUIRED_FIELDS <= set(result)
    assert result["implementationStatus"] == "implemented"
    assert result["builtinMethodKey"] == "llm-analysis"
    assert 0 <= result["threatScore"] <= 100
    assert result["threatLevel"] in {"高", "中", "低"}


def test_target_scores_are_clamped() -> None:
    payload = load_sample()
    payload["targets"][0]["capabilities"]["firepower"] = 99
    payload["targets"][0]["confidence"]["overallConfidence"] = 9
    result = analyze(None, extraction_json=payload, generate_assessment=False)
    for target in result["targetAssessments"]:
        assert 0 <= target["threatScore"] <= 100
        assert 0 <= target["valueScore"] <= 100
        assert 0 <= target["priorityScore"] <= 100


def test_missing_coordinate_does_not_break_output() -> None:
    payload = load_sample()
    payload["targets"][0]["location"]["coordinates"] = None
    payload["targets"][1]["location"]["coordinates"] = None
    result = analyze(None, heatmap_density="low", extraction_json=payload, generate_assessment=False)
    assert result["inputSummary"]["extractedTargetCount"] == 2
    assert result["inputSummary"]["validCoordinateTargetCount"] == 0
    assert result["heatmap"]["grid"] == []
    assert result["visualization"]["entities"] == []


def test_coordinates_remain_lon_lat_alt() -> None:
    result = analyze(None, heatmap_density="low", extraction_json=load_sample(), generate_assessment=False)
    target = result["targetAssessments"][0]
    coordinates = target["location"]["coordinates"]
    assert coordinates == [118.12, 32.04, 0.0]
    assert -180 <= coordinates[0] <= 180
    assert -90 <= coordinates[1] <= 90


@pytest.mark.parametrize(
    ("density", "expected_cells"),
    [
        ("low", 1600),
        ("medium", 6400),
        ("high", 19600),
    ],
)
def test_heatmap_density_grid_sizes(density: str, expected_cells: int) -> None:
    result = analyze(None, heatmap_density=density, extraction_json=load_sample(), generate_assessment=False)
    assert len(result["heatmap"]["grid"]) == expected_cells
    assert len(result["heatmapGeojson"]["features"]) == expected_cells


def test_visualization_entities_have_valid_coordinates_and_radius() -> None:
    result = analyze(None, heatmap_density="low", extraction_json=load_sample(), generate_assessment=False)
    assert result["visualization"]["entities"]
    assert result["heatmapBase64"]
    assert result["targetMapBase64"]
    assert result["combinedMapBase64"]
    for entity in result["visualization"]["entities"]:
        coordinates = entity["coordinates"]
        assert len(coordinates) == 3
        assert -180 <= coordinates[0] <= 180
        assert -90 <= coordinates[1] <= 90
        assert entity["radius"] >= 0


def test_visualization_declares_static_heatmap_image_overlay() -> None:
    result = analyze(None, heatmap_density="low", extraction_json=load_sample(), generate_assessment=False)
    overlays = result["visualization"]["imageOverlays"]
    assert overlays
    overlay = overlays[0]
    assert overlay["id"] == "threat-spatial-field"
    assert overlay["imageBase64Field"] == "heatmapBase64"
    assert overlay["rendering"] == "single-tile-image"
    assert 0.8 <= overlay["alpha"] < 1
    assert overlay["displayVersion"] == "soft-continuous-v2"
    assert overlay["bounds"]["east"] > overlay["bounds"]["west"]
    assert overlay["bounds"]["north"] > overlay["bounds"]["south"]


def test_heatmap_png_has_visible_hotspot_alpha() -> None:
    result = analyze(None, heatmap_density="low", extraction_json=load_sample(), generate_assessment=False)
    image = Image.open(io.BytesIO(base64.b64decode(result["heatmapBase64"]))).convert("RGBA")
    max_alpha = image.getchannel("A").getextrema()[1]
    assert max_alpha >= 200
    assert image.getpixel((0, 0))[3] <= 10
    assert image.getpixel((IMAGE_SIZE // 2, IMAGE_SIZE // 2))[3] > 0


def test_heatmap_png_keeps_zero_threat_background_transparent() -> None:
    heatmap = {
        "statistics": {"maxThreat": 0},
        "bounds": {"minLon": 118.0, "minLat": 31.0, "maxLon": 119.0, "maxLat": 32.0},
        "grid": [
            {"point": {"id": "grid-0-0", "coordinates": [118.0, 31.0, 0]}, "threatScore": 0},
            {"point": {"id": "grid-0-1", "coordinates": [119.0, 31.0, 0]}, "threatScore": 0},
            {"point": {"id": "grid-1-0", "coordinates": [118.0, 32.0, 0]}, "threatScore": 0},
            {"point": {"id": "grid-1-1", "coordinates": [119.0, 32.0, 0]}, "threatScore": 0},
        ],
    }
    image = Image.open(io.BytesIO(base64.b64decode(build_heatmap_base64(heatmap)))).convert("RGBA")
    assert image.getpixel((10, 10))[3] == 0


def test_heatmap_png_softens_cell_boundaries() -> None:
    heatmap = {
        "statistics": {"maxThreat": 100},
        "bounds": {"minLon": 118.0, "minLat": 31.0, "maxLon": 119.0, "maxLat": 32.0},
        "grid": [
            {"point": {"id": "grid-0-0", "coordinates": [118.0, 31.0, 0]}, "threatScore": 0},
            {"point": {"id": "grid-0-1", "coordinates": [119.0, 31.0, 0]}, "threatScore": 0},
            {"point": {"id": "grid-1-0", "coordinates": [118.0, 32.0, 0]}, "threatScore": 100},
            {"point": {"id": "grid-1-1", "coordinates": [119.0, 32.0, 0]}, "threatScore": 0},
        ],
    }
    image = Image.open(io.BytesIO(base64.b64decode(build_heatmap_base64(heatmap)))).convert("RGBA")
    edge_alpha = image.getpixel((IMAGE_SIZE // 2 + 8, IMAGE_SIZE // 4))[3]
    assert edge_alpha > 0


def test_enemy_intentions_fallback_when_llm_omits_them() -> None:
    payload = load_sample()
    payload["globalSituation"]["enemyIntentions"] = []
    payload["globalSituation"]["situationSummary"] = ""
    result = analyze(None, extraction_json=payload, generate_assessment=False)
    assert result["enemyIntentions"]
    assert result["enemyIntentions"][0]["id"] == "intent-fallback-1"
    assert result["enemyIntentions"][0]["description"]


def test_enemy_intentions_preserve_llm_items() -> None:
    payload = load_sample()
    payload["globalSituation"]["enemyIntentions"] = [
        {
            "id": "intent-a",
            "type": "air_defense_denial",
            "name": "防空拒止",
            "description": "依托防空节点形成区域拒止。",
            "confidence": 0.88,
            "evidenceIds": ["ev-2"],
        },
        {
            "id": "intent-b",
            "type": "fire_suppression",
            "name": "火力压制",
            "description": "依托火力群压制主要通道。",
            "confidence": 0.77,
            "evidenceIds": ["ev-1"],
        },
    ]
    result = analyze(None, extraction_json=payload, generate_assessment=False)
    assert [item["id"] for item in result["enemyIntentions"]] == ["intent-a", "intent-b"]
    assert result["enemyIntentions"][0]["score"] == 88.0


def test_assessment_docx_generation_with_mock_json(tmp_path: Path) -> None:
    assessment = {
        "version": 1,
        "enemy_force_type": "未知",
        "enemy_force_type_confidence": 0.5,
        "operational_intent": "敌方可能依托火力、防空和侦察节点形成局部拒止态势。",
        "deployment_posture": "敌方部署呈现火力、防空、侦察和反机降设施分区配置。",
        "threat_summary": "结构化结果显示火力覆盖与防空节点是主要威胁来源。",
        "key_evidence": ["识别到火力覆盖节点", "识别到防空节点"],
        "confidence": 0.72,
    }
    result = analyze(
        None,
        heatmap_density="low",
        extraction_json=load_sample(),
        assessment_json=assessment,
        assessment_output_dir=tmp_path,
    )
    assert result["assessmentReport"]["status"] == "completed"
    assert result["assessmentDocxBase64"]
    docx_path = Path(result["assessmentDocxPath"])
    assert docx_path.exists()
    document = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "敌方作战企图与部署态势研判报告" in text
    assert "一、敌方作战企图" in text
    assert len(document.tables) >= 2


def test_llm_parse_repairs_null_subcategory() -> None:
    payload = load_sample()
    payload["targets"][0]["subCategory"] = None
    parsed = parse_llm_extraction(json.dumps(payload, ensure_ascii=False))
    assert parsed.targets[0].subCategory == ""


def test_llm_parse_repairs_scalar_spatial_context_items() -> None:
    payload = load_sample()
    payload["spatialContext"]["terrain"] = "mixed_hilly_and_open"
    payload["spatialContext"]["weather"] = "clear"
    parsed = parse_llm_extraction(json.dumps(payload, ensure_ascii=False))
    assert parsed.spatialContext.terrain[0]["description"] == "mixed_hilly_and_open"
    assert parsed.spatialContext.weather[0]["description"] == "clear"


def test_llm_parse_repairs_textual_equipment_quantities() -> None:
    payload = load_sample()
    payload["targets"][0]["equipment"] = [
        {"name": "火力装备", "quantity": "estimated", "capabilityTags": "fire"},
        {"name": "防空装备", "quantity": "multiple"},
        {"name": "保障装备", "quantity": "various"},
        {"name": "侦察装备", "quantity": "约 3 套"},
    ]
    parsed = parse_llm_extraction(json.dumps(payload, ensure_ascii=False))
    assert [item.quantity for item in parsed.targets[0].equipment] == [1, 2, 1, 3]
    assert parsed.targets[0].equipment[0].capabilityTags == ["fire"]


def test_llm_parse_does_not_default_point_targets_to_coverage() -> None:
    payload = load_sample()
    payload["targets"][0]["category"] = "command_control"
    payload["targets"][0]["coverage"] = {}
    payload["targets"][1]["category"] = "fire_unit"
    payload["targets"][1]["coverage"] = {}
    parsed = parse_llm_extraction(json.dumps(payload, ensure_ascii=False))
    assert parsed.targets[0].coverage.hasCoverage is False
    assert parsed.targets[0].coverage.radiusMeters == 0
    assert parsed.targets[0].coverage.coverageTypes == []
    assert parsed.targets[1].coverage.hasCoverage is True
    assert parsed.targets[1].coverage.radiusMeters > 0


def test_visualization_omits_coverage_for_point_only_categories() -> None:
    payload = load_sample()
    payload["targets"][0]["category"] = "command_control"
    payload["targets"][0]["coverage"] = {}
    parsed = parse_llm_extraction(json.dumps(payload, ensure_ascii=False))
    result = analyze(None, heatmap_density="low", extraction_json=parsed, generate_assessment=False)
    entity_ids = {item["id"] for item in result["visualization"]["entities"]}
    assert "threat-target-target-001" in entity_ids
    assert "threat-coverage-target-001" not in entity_ids
