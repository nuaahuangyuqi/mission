"""DOCX report generation for second-stage operational assessment."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from shutil import which
from typing import Any, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def generate_operational_assessment_docx(
    *,
    assessment: dict[str, Any],
    structured_output: dict[str, Any],
    output_dir: str | Path,
    file_name: str | None = None,
) -> Path:
    """Generate a DOCX matching the existing operational assessment style."""
    output_root = Path(output_dir)
    output_root.mkdir(parents=True, exist_ok=True)
    target_path = output_root / (file_name or f"operational_assessment_{int(datetime.now().timestamp() * 1000)}.docx")

    doc = Document()
    doc.core_properties.title = "敌方作战企图与部署态势研判报告"
    _configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("敌方作战企图与部署态势研判报告")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(15, 23, 42)

    meta_text = _enemy_force_meta(assessment)
    if meta_text:
        _add_body_paragraph(doc, meta_text)

    _add_heading(doc, "一、敌方作战企图")
    _add_body_paragraph(doc, _text(assessment.get("operational_intent"), "暂无研判。"))

    _add_heading(doc, "二、敌方部署态势")
    _add_body_paragraph(doc, _text(assessment.get("deployment_posture"), "暂无研判。"))

    _add_heading(doc, "三、威胁与火力覆盖摘要")
    _add_body_paragraph(doc, _text(assessment.get("threat_summary"), "暂无摘要。"))

    evidence = [str(item).strip() for item in assessment.get("key_evidence") or [] if str(item).strip()]
    _add_heading(doc, "四、主要依据")
    if evidence:
        for item in evidence[:8]:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(item)
    else:
        _add_body_paragraph(doc, "暂无可归纳依据。")

    _add_heading(doc, "五、高威胁目标")
    _add_high_threat_table(doc, structured_output.get("targetAssessments") or [])

    _add_heading(doc, "六、攻击目标火力覆盖区")
    _add_fire_coverage_table(doc, structured_output.get("fireCoverage") or [])

    doc.save(target_path)
    return target_path


def validate_docx_structure(path: str | Path) -> dict[str, Any]:
    """Lightweight structural QA used when render QA is unavailable."""
    try:
        doc = Document(str(path))
        headings = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return {
            "status": "structural_passed",
            "paragraphCount": len(doc.paragraphs),
            "tableCount": len(doc.tables),
            "hasTitle": any("敌方作战企图与部署态势研判报告" in item for item in headings),
            "hasRequiredTables": len(doc.tables) >= 2,
            "renderStatus": "skipped_soffice_missing" if not (which("soffice") or which("libreoffice")) else "not_run",
        }
    except Exception as exc:
        return {"status": "failed", "error": str(exc)}


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2"):
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(30, 41, 59)


def _add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_heading(text, level=1)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(30, 41, 59)


def _add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(str(text or "").strip())


def _add_high_threat_table(doc: Document, targets: Sequence[dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["编号", "类别", "名称", "威胁度", "价值", "优先级"]
    _fill_header(table.rows[0].cells, headers)
    for target in sorted(targets, key=lambda item: item.get("priorityScore", 0), reverse=True)[:12]:
        row = table.add_row().cells
        row[0].text = str(target.get("id", ""))
        row[1].text = _category_label(target.get("category", ""))
        row[2].text = str(target.get("name", ""))[:36]
        row[3].text = _score(target.get("threatScore"))
        row[4].text = _score(target.get("valueScore"))
        row[5].text = _score(target.get("priorityScore"))


def _add_fire_coverage_table(doc: Document, coverage: Sequence[dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    _fill_header(table.rows[0].cells, ["目标编号", "名称", "半径(km)", "覆盖形态", "威胁度"])
    if not coverage:
        row = table.add_row().cells
        row[0].text = "暂无"
        row[1].text = "未识别火力覆盖区"
        row[2].text = "0.0"
        row[3].text = "无"
        row[4].text = "0.000"
        return
    for item in coverage[:20]:
        row = table.add_row().cells
        row[0].text = str(item.get("sourceUnitId", item.get("id", "")))
        row[1].text = str(item.get("name", ""))[:36]
        row[2].text = f"{float(item.get('coverageKm') or 0):.1f}"
        row[3].text = "全向"
        row[4].text = f"{float(item.get('threatValue') or 0):.3f}"


def _fill_header(cells: Sequence[Any], headers: Sequence[str]) -> None:
    for index, header in enumerate(headers):
        cells[index].text = header
        for paragraph in cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _enemy_force_meta(assessment: dict[str, Any]) -> str:
    force_type = _text(assessment.get("enemy_force_type"), "")
    confidence = assessment.get("enemy_force_type_confidence")
    if force_type:
        try:
            return f"敌方军队类型：{force_type}（置信度 {float(confidence or 0):.0%}）"
        except (TypeError, ValueError):
            return f"敌方军队类型：{force_type}"
    return ""


def _category_label(category: str) -> str:
    return {
        "fire_unit": "火力单位",
        "air_defense": "防空节点",
        "recon_sensor": "侦察预警",
        "command_control": "指挥控制",
        "mobility_unit": "机动兵力",
        "logistics_support": "后勤保障",
        "fortification": "反机降设施",
        "electronic_warfare": "电子对抗",
        "unknown": "未知",
    }.get(category, category or "未知")


def _score(value: Any) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return "0.000"


def _text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip() or default
    if isinstance(value, list):
        return "；".join(str(item).strip() for item in value if str(item).strip()) or default
    if isinstance(value, dict):
        return " ".join(str(item).strip() for item in value.values() if str(item).strip()) or default
    return str(value).strip() or default
