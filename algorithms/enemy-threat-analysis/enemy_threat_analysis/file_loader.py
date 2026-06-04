"""Load user-provided threat analysis documents into plain text."""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from .errors import AnalysisInputError


SUPPORTED_EXTENSIONS = {".txt", ".csv", ".json", ".docx", ".pdf", ".xlsx"}
UNSUPPORTED_EXTENSIONS = {".xls"}


@dataclass(frozen=True)
class LoadedFile:
    file_id: str
    file_name: str
    file_type: str
    path: str
    text: str
    size: int

    def prompt_block(self, max_chars: int = 18_000) -> str:
        content = self.text[:max_chars]
        omitted = max(0, len(self.text) - len(content))
        suffix = f"\n[已截断 {omitted} 个字符]" if omitted else ""
        return (
            f"文件ID: {self.file_id}\n"
            f"文件名: {self.file_name}\n"
            f"文件类型: {self.file_type}\n"
            f"内容:\n{content}{suffix}"
        )


def load_files(file_paths: Iterable[str | Path]) -> list[LoadedFile]:
    paths = [Path(item) for item in file_paths]
    if not paths:
        raise AnalysisInputError("请提供至少一个待分析文件。")

    loaded: list[LoadedFile] = []
    for index, path in enumerate(paths, start=1):
        if not path.exists():
            raise AnalysisInputError(f"文件不存在: {path}")
        if not path.is_file():
            raise AnalysisInputError(f"不是可读取文件: {path}")

        extension = path.suffix.lower()
        if extension in UNSUPPORTED_EXTENSIONS:
            raise AnalysisInputError(f"暂不支持 .xls 文件，请另存为 .xlsx 或 .csv: {path.name}")
        if extension not in SUPPORTED_EXTENSIONS:
            raise AnalysisInputError(f"不支持的文件类型 {extension or '<none>'}: {path.name}")

        text = _read_file_text(path, extension)
        if not text.strip():
            raise AnalysisInputError(f"文件没有可分析文本内容: {path.name}")

        loaded.append(
            LoadedFile(
                file_id=f"file-{index}",
                file_name=path.name,
                file_type=extension,
                path=str(path),
                text=text,
                size=path.stat().st_size,
            )
        )
    return loaded


def _read_file_text(path: Path, extension: str) -> str:
    if extension in {".txt", ".csv", ".json"}:
        text = _read_text(path)
        if extension == ".csv":
            return _csv_to_text(text)
        if extension == ".json":
            return _json_to_text(text)
        return text
    if extension == ".docx":
        return _read_docx(path)
    if extension == ".pdf":
        return _read_pdf(path)
    if extension == ".xlsx":
        return _read_xlsx(path)
    raise AnalysisInputError(f"不支持的文件类型: {extension}")


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


def _csv_to_text(text: str) -> str:
    rows = list(csv.reader(text.splitlines()))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def _json_to_text(text: str) -> str:
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        return text
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - depends on optional dependency
        raise AnalysisInputError("读取 .docx 需要安装 python-docx。") from exc
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join([*paragraphs, *table_rows])


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - depends on optional dependency
        raise AnalysisInputError("读取 .pdf 需要安装 pypdf。") from exc
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except Exception as exc:  # pragma: no cover - depends on optional dependency
        raise AnalysisInputError("读取 .xlsx 需要安装 openpyxl。") from exc
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"工作表: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)
