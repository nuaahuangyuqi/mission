"""Load multiple local force-information files into text bundles."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any


SUPPORTED_SUFFIXES = {".txt", ".csv", ".json", ".docx", ".pdf", ".xlsx"}


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_csv(path: Path) -> str:
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append(dict(row))
    return json.dumps(rows, ensure_ascii=False, indent=2)


def _read_json(path: Path) -> str:
    return json.dumps(json.loads(path.read_text(encoding="utf-8")), ensure_ascii=False, indent=2)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx files") from exc
    doc = Document(path)
    parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to read .pdf files") from exc
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required to read .xlsx files") from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheets: dict[str, list[list[Any]]] = {}
    for sheet in workbook.worksheets:
        rows: list[list[Any]] = []
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None and str(cell).strip() for cell in row):
                rows.append([cell for cell in row])
        sheets[sheet.title] = rows
    return json.dumps(sheets, ensure_ascii=False, indent=2, default=str)


def load_files(files: list[str | Path]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Return successful file bundles and imported-file status records."""

    bundles: list[dict[str, Any]] = []
    imported: list[dict[str, Any]] = []
    for index, file_path in enumerate(files, start=1):
        path = Path(file_path)
        record: dict[str, Any] = {
            "fileId": f"file-{index}",
            "fileName": path.name,
            "fileType": path.suffix.lower(),
            "path": str(path),
            "status": "success",
        }
        try:
            if not path.exists():
                raise FileNotFoundError(f"file does not exist: {path}")
            suffix = path.suffix.lower()
            if suffix == ".xls":
                raise ValueError(".xls is not supported; please convert to .xlsx")
            if suffix not in SUPPORTED_SUFFIXES:
                raise ValueError(f"unsupported file type: {suffix}")
            if suffix == ".txt":
                text = _read_txt(path)
            elif suffix == ".csv":
                text = _read_csv(path)
            elif suffix == ".json":
                text = _read_json(path)
            elif suffix == ".docx":
                text = _read_docx(path)
            elif suffix == ".pdf":
                text = _read_pdf(path)
            elif suffix == ".xlsx":
                text = _read_xlsx(path)
            else:
                text = ""
            bundles.append({**record, "text": text[:120_000]})
            record["characterCount"] = len(text)
        except Exception as exc:  # noqa: BLE001 - keep per-file failures explainable
            record["status"] = "failed"
            record["error"] = str(exc)
        imported.append(record)
    return bundles, imported

