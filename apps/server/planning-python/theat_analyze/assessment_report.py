"""
assessment_report.py — 作战企图与部署态势二次研判报告生成。

流程：
1. 读取原始文档、结构化目标、规则威胁指数、火力覆盖区。
2. 调用大模型输出严格 JSON 研判。
3. 用规则兜底解析 JSON，并生成 .docx 报告。
"""

import json
import logging
import os
import re
import time
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt
from openai import OpenAI

from extractor import parse_file
from schemas import SituationMap

logger = logging.getLogger(__name__)


ASSESSMENT_SYSTEM_PROMPT = """你是一名军事态势研判员。任务：只基于输入材料，生成敌方作战企图与部署态势研判，用于态势展示和DOCX报告。

硬性输出规则：
1. 只输出一个合法 JSON 对象，不要 markdown，不要代码块，不要解释，不要思考过程。
2. 所有字符串必须使用 JSON 字符串，不要换行项目符号；数组元素用短句。
3. 不确定时写“可能/倾向于”，并降低 confidence；不要输出 null。
4. 不要提出我方作战方案、打击建议、路线规划或具体行动步骤。

必须严格使用这个 JSON 结构：
{
  "version": 1,
  "operational_intent": {
    "summary": "120-220字，概括敌方主要企图",
    "judgments": ["判断1，40字以内", "判断2，40字以内", "判断3，40字以内"],
    "confidence": 0.0
  },
  "deployment_posture": {
    "summary": "120-220字，概括敌方部署态势",
    "air_defense": "60-120字，描述防空体系",
    "recon_warning": "60-120字，描述侦察预警",
    "anti_airlanding": "60-120字，描述反机降/阻滞",
    "fire_coverage": "60-120字，描述火力覆盖",
    "confidence": 0.0
  },
  "key_evidence": ["证据1，40字以内", "证据2，40字以内", "证据3，40字以内"],
  "threat_summary": "100-180字，基于规则威胁度和火力覆盖的摘要"
}"""


def _repair_json(raw: str) -> Dict[str, Any]:
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    start = text.find("{")
    end = text.rfind("}")
    if 0 <= start < end:
        text = text[start:end + 1]
    return json.loads(text)


def _as_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        return "；".join(str(item).strip() for item in value if str(item).strip())
    if isinstance(value, dict):
        return "；".join(str(item).strip() for item in value.values() if str(item).strip())
    return str(value).strip()


def _as_confidence(value: Any, default: float = 0.55) -> float:
    try:
        return max(0.0, min(float(value), 1.0))
    except (TypeError, ValueError):
        return default


def _normalize_assessment(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize either the new compact schema or the old flat schema."""
    if "operational_intent" in payload and isinstance(payload.get("operational_intent"), str):
        return {
            "operational_intent": _as_text(payload.get("operational_intent")),
            "deployment_posture": _as_text(payload.get("deployment_posture")),
            "intent_confidence": _as_confidence(payload.get("intent_confidence")),
            "posture_confidence": _as_confidence(payload.get("posture_confidence")),
            "key_evidence": list(payload.get("key_evidence") or [])[:8],
            "threat_summary": _as_text(payload.get("threat_summary")),
        }

    intent = payload.get("operational_intent") or {}
    posture = payload.get("deployment_posture") or {}
    intent_summary = _as_text(intent.get("summary") if isinstance(intent, dict) else intent)
    judgments = intent.get("judgments") if isinstance(intent, dict) else []
    if judgments:
        intent_summary = f"{intent_summary} " + "；".join(str(item).strip() for item in judgments if str(item).strip())

    posture_parts = []
    if isinstance(posture, dict):
        for key in ("summary", "air_defense", "recon_warning", "anti_airlanding", "fire_coverage"):
            text = _as_text(posture.get(key))
            if text:
                posture_parts.append(text)
    else:
        posture_parts.append(_as_text(posture))

    return {
        "operational_intent": intent_summary,
        "deployment_posture": " ".join(posture_parts).strip(),
        "intent_confidence": _as_confidence(intent.get("confidence") if isinstance(intent, dict) else None),
        "posture_confidence": _as_confidence(posture.get("confidence") if isinstance(posture, dict) else None),
        "key_evidence": [str(item).strip() for item in (payload.get("key_evidence") or []) if str(item).strip()][:8],
        "threat_summary": _as_text(payload.get("threat_summary")),
    }


def _fallback_assessment(
    situation: SituationMap,
    targets_data: List[Dict[str, Any]],
    fire_coverage_areas: List[Dict[str, Any]],
) -> Dict[str, Any]:
    categories: Dict[str, int] = {}
    for target in targets_data:
        category = target.get("target_category") or "未分类"
        categories[category] = categories.get(category, 0) + 1

    enemy_type = situation.enemy_force_type or "未知"
    high_threat = sorted(targets_data, key=lambda item: item.get("threat_index", 0), reverse=True)[:5]
    high_names = "、".join(t.get("target_id", "") for t in high_threat if t.get("target_id")) or "未形成明确高威胁目标"
    category_text = "，".join(f"{key}{value}处" for key, value in sorted(categories.items(), key=lambda item: item[1], reverse=True)[:6])

    return {
        "operational_intent": (
            f"综合文档线索和结构化目标，敌方类型判断为{enemy_type}。目标体系以{category_text}为主，"
            f"高威胁目标集中在{high_names}。从直升机编组威胁视角看，敌方企图更偏向通过低空预警、"
            "防空拦截和反机降阻滞形成纵深拒止，压缩直升机突入、悬停、卸载和撤离窗口。"
        ),
        "deployment_posture": (
            "部署态势呈现节点化防御特征：防空阵地和雷达预警节点负责低空发现与拦截，"
            "火力节点依据射程形成覆盖区，步兵阵地、预备队和阻滞设施承担反机降封控。"
            f"当前共识别{len(fire_coverage_areas)}个攻击目标覆盖区，应视为火力与防空影响范围的基础边界。"
        ),
        "intent_confidence": 0.55,
        "posture_confidence": 0.55,
        "key_evidence": [
            f"目标类别分布：{category_text or '暂无'}",
            f"高威胁目标：{high_names}",
            f"火力覆盖区数量：{len(fire_coverage_areas)}",
        ],
        "threat_summary": (
            "规则算法已按直升机编组场景提高防空拦截、侦察预警和反机降阻滞权重，"
            "火力覆盖范围由所有攻击目标的攻击半径生成，用于支撑态势研判和地图叠加。"
        ),
    }


def _summarize_targets(targets_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    summarized = []
    for target in sorted(targets_data, key=lambda item: item.get("threat_index", 0), reverse=True)[:80]:
        factors = target.get("factors") or {}
        summarized.append({
            "id": target.get("target_id"),
            "category": target.get("target_category"),
            "name": target.get("target_name"),
            "threat_index": round(float(target.get("threat_index", 0)), 4),
            "lat": target.get("lat"),
            "lng": target.get("lng"),
            "air_defense_score": factors.get("air_defense_score"),
            "recon_warning_score": factors.get("recon_warning_score"),
            "anti_airlanding_score": factors.get("anti_airlanding_score"),
            "range_km": factors.get("lethality_range_km"),
            "description": target.get("description", "")[:120],
        })
    return summarized


def _call_llm_assessment(
    doc_paths: List[str],
    situation: SituationMap,
    targets_data: List[Dict[str, Any]],
    fire_coverage_areas: List[Dict[str, Any]],
    model: str,
    api_key: str,
    base_url: str,
) -> Dict[str, Any]:
    documents = []
    for idx, path in enumerate(doc_paths):
        content = parse_file(path)
        documents.append({
            "filename": os.path.basename(path),
            "content_excerpt": content[:2500],
        })

    payload = {
        "enemy_force_type": situation.enemy_force_type,
        "enemy_force_type_confidence": situation.enemy_force_type_confidence,
        "enemy_force_type_basis": situation.enemy_force_type_basis,
        "targets": _summarize_targets(targets_data)[:40],
        "fire_coverage_areas": fire_coverage_areas[:60],
        "documents": documents,
    }

    client = OpenAI(api_key=api_key, base_url=base_url)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": ASSESSMENT_SYSTEM_PROMPT},
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
        ],
        response_format={"type": "json_object"},
        temperature=0.2,
        max_tokens=2048,
    )
    content = response.choices[0].message.content or "{}"
    return _normalize_assessment(_repair_json(content))


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"


def _add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(18)
    paragraph.paragraph_format.line_spacing = 1.25


def _write_docx(
    output_path: Path,
    situation: SituationMap,
    assessment: Dict[str, Any],
    targets_data: List[Dict[str, Any]],
    fire_coverage_areas: List[Dict[str, Any]],
) -> None:
    doc = Document()
    doc.core_properties.title = "敌方作战企图与部署态势研判报告"

    _add_heading(doc, "敌方作战企图与部署态势研判报告", 0)
    _add_paragraph(doc, f"敌方军队类型：{situation.enemy_force_type}（置信度 {situation.enemy_force_type_confidence:.0%}）")
    if situation.enemy_force_type_basis:
        _add_paragraph(doc, f"类型判断依据：{situation.enemy_force_type_basis}")

    _add_heading(doc, "一、敌方作战企图", 1)
    _add_paragraph(doc, assessment.get("operational_intent", "暂无研判。"))

    _add_heading(doc, "二、敌方部署态势", 1)
    _add_paragraph(doc, assessment.get("deployment_posture", "暂无研判。"))

    _add_heading(doc, "三、威胁与火力覆盖摘要", 1)
    _add_paragraph(doc, assessment.get("threat_summary", "暂无摘要。"))

    evidence = assessment.get("key_evidence") or []
    if evidence:
        _add_heading(doc, "四、主要依据", 1)
        for item in evidence[:8]:
            doc.add_paragraph(str(item), style="List Bullet")

    _add_heading(doc, "五、高威胁目标", 1)
    top_targets = sorted(targets_data, key=lambda item: item.get("threat_index", 0), reverse=True)[:12]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["编号", "类别", "名称", "威胁度", "防空", "预警/反机降"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for target in top_targets:
        factors = target.get("factors") or {}
        row = table.add_row().cells
        row[0].text = str(target.get("target_id", ""))
        row[1].text = str(target.get("target_category", ""))
        row[2].text = str(target.get("target_name", ""))[:30]
        row[3].text = f"{float(target.get('threat_index', 0)):.3f}"
        row[4].text = str(factors.get("air_defense_score", ""))
        row[5].text = f"{factors.get('recon_warning_score', '')}/{factors.get('anti_airlanding_score', '')}"

    _add_heading(doc, "六、攻击目标火力覆盖区", 1)
    coverage_table = doc.add_table(rows=1, cols=5)
    coverage_table.style = "Table Grid"
    for idx, header in enumerate(["目标编号", "类别", "半径(km)", "覆盖形态", "威胁度"]):
        coverage_table.rows[0].cells[idx].text = header
    for area in fire_coverage_areas[:20]:
        row = coverage_table.add_row().cells
        row[0].text = str(area.get("target_id", ""))
        row[1].text = str(area.get("target_category", ""))
        row[2].text = f"{float(area.get('radius_km', 0)):.1f}"
        row[3].text = "扇区" if area.get("coverage_type") == "sector" else "全向"
        row[4].text = f"{float(area.get('threat_index', 0)):.3f}"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_operational_assessment_docx(
    doc_paths: List[str],
    situation: SituationMap,
    targets_data: List[Dict[str, Any]],
    fire_coverage_areas: List[Dict[str, Any]],
    output_dir: str,
    model: str,
    api_key: str,
    base_url: str,
) -> Dict[str, Any]:
    """生成二次研判 JSON 与 docx 文件。"""
    try:
        assessment = _call_llm_assessment(
            doc_paths=doc_paths,
            situation=situation,
            targets_data=targets_data,
            fire_coverage_areas=fire_coverage_areas,
            model=model,
            api_key=api_key,
            base_url=base_url,
        )
    except Exception as exc:
        logger.warning("二次研判 LLM 调用失败，使用规则兜底: %s", exc)
        assessment = _fallback_assessment(situation, targets_data, fire_coverage_areas)

    output_path = Path(output_dir) / f"operational_assessment_{int(time.time() * 1000)}.docx"
    _write_docx(output_path, situation, assessment, targets_data, fire_coverage_areas)

    return {
        "assessment": assessment,
        "docx_path": str(output_path),
    }
