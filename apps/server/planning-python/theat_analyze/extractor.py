"""
extractor.py — 文档解析与 LLM 结构化数据提取模块

本模块负责：
1. 使用 python-docx 解析 .docx 文档，提取正文段落与表格内容。
2. 支持两种 LLM 后端：
   - Ollama 本地模型 (通过 ollama 库)
   - 外部 API (通过 openai 库，兼容 OpenAI / DeepSeek / 其他兼容网关)
3. 强制要求 LLM 提取，无 LLM 配置时直接报错。
4. 包含最多 3 次重试的容错逻辑。
"""

import json
import logging
import os
import re
import sys
import time
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import List, Optional

from docx import Document

from schemas import SituationMap

logger = logging.getLogger(__name__)


# ────────────────────────────────────────────────────────
# LLM 后端枚举
# ────────────────────────────────────────────────────────

class LLMBackend(str, Enum):
    """LLM 后端类型枚举。"""
    OLLAMA = "ollama"
    OPENAI_API = "openai_api"


@dataclass(frozen=True)
class GenerationBudget:
    """LLM generation limits resolved from document size and env overrides."""

    document_chars: int
    context_tokens: int
    output_tokens: int
    profile_name: str


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None or raw.strip() == "":
        return default
    try:
        value = int(raw)
    except ValueError:
        logger.warning("环境变量 %s=%r 不是整数，使用默认值 %d", name, raw, default)
        return default
    return max(value, 1)


def _resolve_generation_budget(master_content: str) -> GenerationBudget:
    """Resolve context/output sizes for 10k-100k Chinese-character documents.

    Defaults are tuned for extracting roughly 50-200 targets. They can be
    overridden without editing code:
      - LLM_CONTEXT_TOKENS: Ollama num_ctx / expected model context
      - LLM_OUTPUT_TOKENS: Ollama num_predict / API max_tokens
    """
    document_chars = len(re.sub(r"\s+", "", master_content))

    if document_chars <= 20_000:
        profile_name = "short_10k_20k"
        default_context = 32_768
        default_output = 16_384
    elif document_chars <= 50_000:
        profile_name = "medium_20k_50k"
        default_context = 65_536
        default_output = 24_576
    else:
        profile_name = "long_50k_100k"
        default_context = 131_072
        default_output = 32_768

    return GenerationBudget(
        document_chars=document_chars,
        context_tokens=_env_int("LLM_CONTEXT_TOKENS", default_context),
        output_tokens=_env_int("LLM_OUTPUT_TOKENS", default_output),
        profile_name=profile_name,
    )


def _resolve_think_mode() -> bool:
    """Return whether model thinking mode should be requested in prompts."""
    raw = os.getenv("LLM_THINK_MODE", "off").strip().lower()
    return raw in {"1", "true", "yes", "on", "think", "thinking", "enable", "enabled"}


def _build_think_instruction(think_enabled: bool) -> str:
    if think_enabled:
        return (
            "Think 模式：开启。请使用 /think 进行内部推理，先在内部检查目标是否遗漏、"
            "坐标是否可解析、JSON 是否闭合；最终答案仍必须只保留 JSON。"
        )
    return (
        "Think 模式：关闭。请使用 /no_think，禁止输出 `<think>`、`</think>` 或任何推理过程，"
        "直接输出最终 JSON。"
    )


def _build_budget_instruction(budget: GenerationBudget) -> str:
    return (
        f"文档规模：约 {budget.document_chars} 个非空白字符；"
        f"预算档位：{budget.profile_name}；"
        f"上下文预算：{budget.context_tokens} tokens；"
        f"输出预算：{budget.output_tokens} tokens。"
        "预计输出目标数量为 50-200 个。若目标数量接近 200 个，"
        "description 压缩到 50-80 字，优先保证字段完整、坐标可解析、JSON 合法闭合。"
    )


def _field(obj, name: str, default=None):
    """Read a field from dict-like or object-like streaming chunks."""
    if isinstance(obj, dict):
        return obj.get(name, default)
    return getattr(obj, name, default)


def _extract_ollama_stream_text(chunk) -> tuple[str, str]:
    """Return (content, thinking) from an Ollama streaming chunk."""
    message = _field(chunk, "message", {}) or {}
    content = _field(message, "content", "") or ""
    thinking = (
        _field(message, "thinking", "")
        or _field(chunk, "thinking", "")
        or ""
    )
    return str(content), str(thinking)


def _extract_openai_stream_text(delta) -> tuple[str, str]:
    """Return (content, reasoning) from OpenAI-compatible streaming deltas."""
    if delta is None:
        return "", ""
    content = _field(delta, "content", "") or ""
    reasoning = (
        _field(delta, "reasoning_content", "")
        or _field(delta, "reasoning", "")
        or _field(delta, "thinking", "")
        or _field(delta, "thinking_content", "")
        or ""
    )
    return str(content), str(reasoning)


# ────────────────────────────────────────────────────────
# 文档解析
# ────────────────────────────────────────────────────────

def parse_file(file_path: str) -> str:
    """解析 .docx 或 .txt 文件，将内容提取为纯文本字符串。

    对于 .docx，会将段落与表格内容拼接，表格使用 Markdown 风格格式化。
    对于 .txt，直接读取全部文本。

    Args:
        file_path: 文件的绝对或相对路径。

    Returns:
        包含所有文本内容的纯文本字符串。
    """
    if str(file_path).endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
            
    doc = Document(file_path)
    parts: List[str] = []

    # 提取所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 提取所有表格（Markdown 格式化，LLM 更易理解）
    for idx, table in enumerate(doc.tables):
        parts.append(f"\n===== 表格 {idx + 1} =====")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
            # 在表头行后插入分隔线
            if row_idx == 0:
                parts.append(" | ".join(["---"] * len(cells)))

    return "\n".join(parts)


# ────────────────────────────────────────────────────────
# 优化后的提示词
# ────────────────────────────────────────────────────────

SYSTEM_PROMPT = """# 角色
你是一名资深军事情报结构化数据提取专家 (SIGINT/GEOINT 分析员)。
你的任务是：从用户提供的军事情报文档与装备手册中，提取所有可标定或可估算坐标的目标实体，输出严格合法的 JSON。

# 输出格式
你只能输出一个 JSON 对象，顶层键包含 "enemy_force_type"、"enemy_force_type_confidence"、"enemy_force_type_basis"、"targets"。
不要输出解释、分析过程、markdown 代码块、注释、前后缀文字。

## 顶层字段

### enemy_force_type (string, 必填)
基于文档内容和你的世界知识判断敌方军队类型。
可取值示例：台军 / 越军 / 美军 / 英军 / 日军 / 韩军 / 菲律宾军 / 印军 / 俄军 / 其他 / 未知。
如果文档包含装备名称、编制番号、地名、战术术语或军种体系特征，必须综合判断，不要只按国名关键词机械匹配。

### enemy_force_type_confidence (float, 必填)
军队类型判断置信度，0.0-1.0。

### enemy_force_type_basis (string, 必填)
50-150字说明判断依据，引用文档中的装备、编制、地理或术语线索，并可结合世界知识。

每个目标实体必须包含以下字段，字段名不能改，不能多包一层对象：

## 字段定义

### target_id (string, 必填)
目标编号。直接从文档的"数量/编号"列提取，取 "/" 后的编号部分。
示例: "C2-01", "SAM-01", "EA-01", "SHORAD-02"

### target_category (string, 必填)
目标类别。直接从文档的"目标类别"列提取原文。
可能的值: C2指挥 / 雷达预警 / 防空阵地 / 炮兵阵地 / 电子战 / 步兵阵地 / 后勤节点 / 后勤/支撑 / 火力节点 / 通信节点 / 侦察节点 / 预备队 / 预警节点 / 电子战/诱饵 / 火力坐标

### target_name (string, 必填)
目标名称。取自文档的"目标名称"列完整文本。

### description (string, 必填)
根据情报原文提取并浓缩的目标战术描述（50-150字）。
**必须包含**：部署位置特征、所承担的战术职能、威胁特点。
若原文描述详尽，直接摘要原文；若原文简略，结合军事常识进行战术解读与推断。
示例: "该阵地部署于山坡反斜面阵地，主要承担防空拦截任务，配备猎鹰-9中程防空导弹，具备对中低空目标的持续拦截能力，与后方C2节点通过光纤链接。"

### raw_coordinates (string, 必填)
目标坐标。该字段会被下游程序直接解析，**必须是可解析的 WGS84 经纬度字符串**。

**允许格式，只使用下面两类之一**:
- 十进制度: "23.2885°N, 114.0078°E"
- 裸十进制度: "23.2885, 114.0078"

**坐标缺失时的处理规则（非常重要）**:
1. 如果原文有经纬度坐标，必须原样复制，不得修改、四舍五入或改写。
2. 如果原文没有经纬度，但有地名/驻地/行政区/设施名，必须根据地名估算该地点中心点或设施近似位置，并输出十进制度坐标。
   示例: 原文只写"位于新北市淡水区"时，可输出 "25.1710°N, 121.4430°E"，同时把 confidence 降到 0.3-0.6。
3. 如果只有上级区域，没有精确地名，使用该区域的行政中心或战区中心点作为近似坐标，并把 confidence 降到 0.1-0.4。
4. 如果完全无法确认地点，也必须选择文档中主战区的中心点作为临时坐标，并在 description 中说明该坐标为估算；**绝不能输出占位值**。

**严禁输出以下任何值**:
"N/A"、"N\\A"、"NA"、"无"、"未知"、"不详"、"未提供"、"null"、"None"、"-"、空字符串、纯地名。

**自检规则**: 每个 raw_coordinates 必须同时满足：
- 至少包含两个阿拉伯数字；
- 纬度和经度之间必须有英文逗号或中文逗号；
- 不能是占位值；
- 不能是对象、数组或单独的 lat/lng 字段，必须是字符串。

**错误示例**:
- "raw_coordinates": "N/A"
- "raw_coordinates": "新北市淡水区"
- "raw_coordinates": {"lat": 25.171, "lng": 121.443}

**正确示例**:
- "raw_coordinates": "25.1710°N, 121.4430°E"
- "raw_coordinates": "25.1710, 121.4430"

### heading_angle (float, 必填)
主要威胁朝向角 (0-360, 正北=0, 顺时针递增)。

**判断规则（按优先级从高到低）**:
1. 全向设备 → 设为 -1:
   - 全向旋转搜索雷达 (如预警雷达 360° 扫描)
   - 全向射界火炮 (如 "雷神-M" 迫榴炮, 360° 方向射界)
   - 被动传感器/测向站 (无辐射源)
   - 指挥所 (C2)、后勤节点、通信节点、预备队
   - 全向微波穹顶 (如 "天幕" 系统, 扇区扫描覆盖全域)
   - 诱饵/假目标 (全向辐射模拟)
   - 观察所、侦察哨、无人机场

2. 有方向设备 → 根据文档上下文判断朝向:
   - "迎敌面" / 面向南方 → 约 180°-210°
   - 防空导弹阵地 (反斜面部署, 向南/西南方向拦截) → 约 200°-225°
   - 火控雷达 (照射方向) → 约 200°-225°
   - 定向干扰天线 → 根据 "迎敌面" 判断, 约 210°
   - 反装甲伏击阵位 (沿公路方向) → 约 180°-200°
   - 便携防空导弹 → 根据具体部署位置判断射界方向

### factors (object, 必填)
威胁量化因子，包含 7 个子字段。下游威胁识别模型面向我方直升机编组，必须优先刻画敌方防空体系、侦察预警和反机降设施。

#### lethality_range_km (float, 必填)
最大硬杀伤或软杀伤有效距离 (km)。
**必须结合装备手册的性能参数表确定**:
- "猎鹰-9" 防空导弹: 杀伤区远界 50 km
- "蜂刺-X" 便携防空导弹: 有效射程 6.5 km
- "极光" 电战模块: ERP 覆盖 60 km
- "天幕" 微波穹顶: 有效致盲距离 8 km
- "雷神-M" 迫榴炮: 最大射程 15 km
- S波段预警雷达: 探测距离 65 km
- 反坦克导弹: 约 4 km
- 被动测向站: 监测距离约 20 km
- 指挥所/后勤/通信/侦察: 0-1 km (无直接杀伤力)
- 诱饵/假目标: 辐射覆盖半径约 2-5 km

#### ew_erp_mw (float, 必填)
电子战有效辐射功率 (MW)。
- "极光" 电战模块: > 1.0 MW
- "天幕" 微波穹顶: 约 0.5 MW (脉冲平均)
- 诱饵辐射源: 约 0.01 MW
- 前置中继天线: 约 0.3 MW
- **所有非电战设备: 必须设为 0.0**

#### survivability_score (int, 1-10, 必填)
生存能力评分:
- 9-10: 深层地下坑道 (如 Alpha 深层堡垒, 25m+ 岩层覆盖)
- 7-8: 反斜面掩体、加固溶洞、混凝土塔基
- 5-6: 伪装掩体、林地隐蔽、密林部署
- 3-4: 浅层工事、标准野战工事
- 1-2: 暴露阵位、充气假目标、开阔地

#### target_value (int, 1-10, 必填)
目标高价值指数:
- 9-10: 旅级 C2 指挥所、主预警雷达、光纤交换中心
- 7-8: 防空导弹连、电战主模块、火控雷达、主发电机房
- 5-6: 营级指挥所、弹药库、后勤枢纽、通信中继
- 3-4: 侦察哨、预备队、前推囤积点
- 1-2: 诱饵假目标、普通步兵哨

#### air_defense_score (int, 1-10, 必填)
对直升机编组的防空拦截威胁评分：
- 9-10: 中远程防空导弹、综合防空指挥/火控节点、具备多目标拦截能力的防空阵地
- 7-8: 近程防空导弹、高炮/弹炮合一、伴随野战防空分队
- 5-6: 便携防空、重机枪/小口径高炮、可对低空慢速目标射击的步兵火力
- 1-4: 非防空目标，或仅具有限制性低空自卫火力

#### recon_warning_score (int, 1-10, 必填)
侦察预警与低空发现能力评分：
- 9-10: 主预警雷达、低空补盲雷达、被动测向/电子侦察网、综合预警节点
- 7-8: 火控雷达、无人机侦察节点、前沿观察/传感器组网
- 4-6: 普通观察哨、通信中继、可提供间接预警的信息节点
- 1-3: 无明显侦察预警能力

#### anti_airlanding_score (int, 1-10, 必填)
反机降/反直升机着陆阻滞能力评分：
- 9-10: 机降场障碍、反机降火力伏击区、预设爆破/拒止设施、机场/开阔地反着陆部署
- 7-8: 步兵反机降阵地、预备队快速反击节点、道路/桥梁封控与机动阻滞
- 4-6: 可对低空降落和卸载地域形成压制的炮兵/迫榴炮/机枪火力
- 1-3: 与反机降关系较弱的支撑目标

### equip_params (object, 可选)
其他关键装备参数字典。从装备手册中提取关键性能指标。
示例: {"max_targets": "8", "reaction_time_s": "6-8", "pk": "0.85"}

### confidence (float, 必填)
数据置信度 (0.0 到 1.0)。
- 原文明确写出的给 0.9-1.0
- 部分缺失靠常识推断的给 0.6-0.8
- 完全缺失靠纯猜测补全的给 0.1-0.5

# 关键约束
1. **绝不能遗漏目标**。提取文档中提及的所有具体军事目标。
2. **所有字段无条件补全**。如果原文缺失坐标、装备参数、部署状态、甚至具体的战术数据，必须用地名、上下文和军事常识估算，绝不能报错、跳过或返回空值。
3. **如实打分置信度**。根据你脑补的程度，如实填写 `confidence` 字段。
4. **必须结合装备手册的定量数据**填写 factors，若手册也未提及，依靠猜测填写并调低置信度。
5. **严禁破坏 JSON 格式**。输出必须是完全合法的纯 JSON 文本，不能包含任何 markdown 代码块标记 (如 ```json) 或前后思考文字！
6. 仔细检查文档中所有表格，不要遗漏任何有坐标的目标行。
7. 输出前逐个目标检查 `raw_coordinates`，发现 N/A/未知/纯地名/缺少逗号/少于两个数字时，必须立刻改成估算经纬度。
8. 输出前逐个目标检查 factors：七个数值字段必须存在，非电战设备的 ew_erp_mw 必须为 0.0。
9. 如果不确定某项信息，填入合理估算值并降低 confidence，不要输出“未知”。"""

USER_PROMPT_TEMPLATE = """请从以下军事情报文档合集中提取所有目标实体的结构化数据。

合集中可能包含多份文件（如敌情态势分析、目标部署表、装备手册等）。你需要综合这些情报，提取所有具体目标。

你需要：
1. 提取文档中提及的所有目标（尤其是表格中的部署节点）。
2. 先判断敌方军队类型，再结合文档中提及的性能参数来精确填写 factors 字段。若无参数或装备手册缺失，必须根据军事常识进行大胆预测。
3. 根据目标部署描述和地形上下文判断 heading_angle。
4. 为每个目标撰写 description 字段，浓缩原文关键战术信息（50-150字）。
5. 对每个目标输出可解析的 raw_coordinates；没有原始坐标时，根据地名估算近似经纬度，严禁使用 N/A/未知/null 等占位值。
6. 输出前只做内部检查，不要把检查过程写出来。
7. 严格遵守下面的 Think 模式开关。

================== 生成预算与模式 ==================
{budget_instruction}

{think_instruction}

必须使用以下 JSON 骨架，字段名不要改：
{
  "enemy_force_type": "台军",
  "enemy_force_type_confidence": 0.85,
  "enemy_force_type_basis": "根据文档中的装备、编制和地名线索判断敌方军队类型。",
  "targets": [
    {
      "target_id": "C2-01",
      "target_category": "C2指挥",
      "target_name": "目标名称",
      "description": "50-150字战术描述",
      "raw_coordinates": "23.2885°N, 114.0078°E",
      "heading_angle": -1,
      "factors": {
        "lethality_range_km": 0.0,
        "ew_erp_mw": 0.0,
        "survivability_score": 8,
        "target_value": 10,
        "air_defense_score": 2,
        "recon_warning_score": 6,
        "anti_airlanding_score": 3
      },
      "equip_params": {},
      "confidence": 0.8
    }
  ]
}

================== 情报合集原文 ==================
{all_documents_content}

请输出完整的纯 JSON 结果，确保不遗漏任何目标，大胆推断缺失数据并给出 confidence 评分。输出前必须检查所有 raw_coordinates 都是坐标字符串。严禁任何思考文字或 markdown 代码块语法！"""


def _render_user_prompt(
    master_content: str,
    budget: GenerationBudget,
    think_enabled: bool,
) -> str:
    """Render the user prompt without treating JSON example braces as format fields."""
    return (
        USER_PROMPT_TEMPLATE
        .replace("{budget_instruction}", _build_budget_instruction(budget))
        .replace("{think_instruction}", _build_think_instruction(think_enabled))
        .replace("{all_documents_content}", master_content)
    )


# ────────────────────────────────────────────────────────
# JSON 修复工具
# ────────────────────────────────────────────────────────

def _repair_json(raw: str) -> str:
    """尝试从 LLM 原始输出中提取并修复 JSON 字符串。

    处理常见问题：markdown 代码块包裹、前后多余文本等。

    Args:
        raw: LLM 的原始输出文本。

    Returns:
        清理后的 JSON 字符串。
    """
    text = raw.strip()
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()

    # 去除 markdown 代码块包裹
    if text.startswith("```"):
        # 去掉首行 (```json 或 ```)
        text = re.sub(r"^```(?:json)?\s*\n?", "", text)
        # 去掉末尾 ```
        text = re.sub(r"\n?```\s*$", "", text)

    import json
    # 尝试整体解析
    try:
        json.loads(text)
        return text
    except json.JSONDecodeError:
        pass

    # 尝试从前后多余文本中截取 JSON 对象或数组主体。
    object_start = text.find("{")
    object_end = text.rfind("}")
    if 0 <= object_start < object_end:
        candidate = text[object_start:object_end + 1]
        try:
            json.loads(candidate)
            return candidate
        except json.JSONDecodeError:
            pass

    array_start = text.find("[")
    array_end = text.rfind("]")
    if 0 <= array_start < array_end:
        candidate = text[array_start:array_end + 1]
        try:
            json.loads(candidate)
            return candidate
        except json.JSONDecodeError:
            pass

    # 处理截断的 JSON (用户暂停推演或生成中断)
    # 倒序查找，尝试在每一个 '}' 处截断并闭合，直到解析成功
    idx = len(text) - 1
    while idx >= 0:
        idx = text.rfind('}', 0, idx + 1)
        if idx == -1:
            break
            
        truncated = text[:idx + 1]
        
        # 尝试几种常见的补全模式
        closures = ["", "]", "]}", "]}", "}"]
        for closure in closures:
            try:
                candidate = truncated + closure
                json.loads(candidate)
                return candidate
            except json.JSONDecodeError:
                continue
                
        idx -= 1  # 没成功，继续往前找上一个 '}'
        
    return '{"targets": []}'


def _coordinate_pair_to_string(first, second) -> Optional[str]:
    """Convert a numeric lat/lon pair to the parser-friendly coordinate string."""
    try:
        lat = float(first)
        lon = float(second)
    except (TypeError, ValueError):
        return None

    # Some weaker models emit lon,lat even when asked for lat,lon.
    if abs(lat) > 90 and abs(lon) <= 90:
        lat, lon = lon, lat

    if not (-90 <= lat <= 90 and -180 <= lon <= 180):
        return None

    return f"{lat:.6f}, {lon:.6f}"


def _coordinate_from_mapping(value: dict) -> Optional[str]:
    lat = next(
        (value[key] for key in ("lat", "latitude", "纬度") if key in value and value[key] is not None),
        None,
    )
    lon = next(
        (value[key] for key in ("lng", "lon", "longitude", "经度") if key in value and value[key] is not None),
        None,
    )
    return _coordinate_pair_to_string(lat, lon)


def _is_coordinate_placeholder(value) -> bool:
    text = str(value).strip().lower().replace("\\", "/")
    text = re.sub(r"\s+", "", text)
    return text in {
        "",
        "n/a",
        "na",
        "none",
        "null",
        "unknown",
        "unk",
        "无",
        "未知",
        "不详",
        "未提供",
        "待定",
        "-",
    }


def _normalize_coordinate_shapes(payload: dict) -> dict:
    """Normalize common weak-model coordinate shapes before Pydantic validation."""
    targets = payload.get("targets")
    if not isinstance(targets, list):
        return payload

    for target in targets:
        if not isinstance(target, dict):
            continue

        raw = target.get("raw_coordinates")
        normalized: Optional[str] = None

        if isinstance(raw, dict):
            normalized = _coordinate_from_mapping(raw)
        elif isinstance(raw, (list, tuple)) and len(raw) >= 2:
            normalized = _coordinate_pair_to_string(raw[0], raw[1])

        if not normalized and _is_coordinate_placeholder(raw):
            normalized = _coordinate_from_mapping(target)

        if normalized:
            target["raw_coordinates"] = normalized

    return payload


def _validate_situation_coordinates(situation: SituationMap) -> None:
    """Fail fast if LLM coordinates cannot be parsed by the downstream engine."""
    from geo_math import parse_coordinate_string

    invalid_targets = []
    for target in situation.targets:
        try:
            parse_coordinate_string(target.raw_coordinates)
        except Exception as exc:
            target_id = target.target_id or "<unknown>"
            invalid_targets.append(
                f"{target_id} raw_coordinates={target.raw_coordinates!r} ({exc})"
            )

    if invalid_targets:
        preview = "; ".join(invalid_targets[:8])
        extra = "" if len(invalid_targets) <= 8 else f"; 另有 {len(invalid_targets) - 8} 个目标"
        raise ValueError(
            "LLM 输出了下游无法解析的坐标，请把这些 raw_coordinates 改成十进制度坐标: "
            f"{preview}{extra}"
        )


def _parse_and_validate_situation(raw_json: str) -> SituationMap:
    """Repair raw LLM JSON, normalize common shapes, and validate coordinates."""
    cleaned = _repair_json(raw_json)
    payload = json.loads(cleaned)
    if isinstance(payload, list):
        payload = {"targets": payload}
    if not isinstance(payload, dict):
        raise ValueError("LLM 输出必须是 JSON 对象，顶层键为 targets")
    payload = _normalize_coordinate_shapes(payload)
    situation = SituationMap.model_validate(payload)
    if not situation.targets:
        raise ValueError("LLM 输出中 targets 为空，请重新提取文档中的具体目标实体")
    _validate_situation_coordinates(situation)
    return situation


def _build_attempt_user_msg(base_user_msg: str, last_error: Optional[Exception]) -> str:
    """Add a focused repair instruction after a failed extraction attempt."""
    if not last_error:
        return base_user_msg

    error_text = str(last_error)
    if len(error_text) > 1200:
        error_text = error_text[:1200] + "..."

    return f"""{base_user_msg}

================== 上一次输出未通过程序校验 ==================
错误信息:
{error_text}

请重新输出完整 JSON。重点修正：
1. 每个 raw_coordinates 都必须是字符串，格式类似 "23.2885°N, 114.0078°E" 或 "23.2885, 114.0078"。
2. 绝对不要输出 N/A、N\\A、NA、未知、无、null、None、-、空字符串或纯地名。
3. 原文没有经纬度时，根据目标的地名、驻地、行政区或战区上下文估算近似经纬度，并降低 confidence。
4. 不要解释，不要输出 markdown，只输出完整纯 JSON。"""


# ────────────────────────────────────────────────────────
# Ollama 本地模型提取
# ────────────────────────────────────────────────────────

def extract_with_ollama(
    doc_paths: List[str],
    model: str = "qwen2.5:32b",
    ollama_host: Optional[str] = None,
    max_retries: int = 3,
) -> SituationMap:
    """使用 Ollama 本地模型从文档中提取结构化态势数据。

    通过 ollama 库调用本地部署的大模型，使用 format 参数约束
    输出为符合 SituationMap JSON Schema 的结构化数据。

    Args:
        doc_paths: .docx 文件路径列表（第一个为情报文档，第二个为装备手册）。
        model: Ollama 模型名称（如 "qwen2.5:32b", "llama3:70b"）。
        ollama_host: Ollama 服务地址（如 "http://localhost:11434"），
                     None 则使用默认地址。
        max_retries: 最大重试次数。

    Returns:
        SituationMap 实例，包含所有提取的目标实体。

    Raises:
        RuntimeError: 超过最大重试次数后仍解析失败。
        ConnectionError: 无法连接到 Ollama 服务。
    """
    import ollama

    # 解析所有文档
    all_content = []
    import os
    for idx, path in enumerate(doc_paths):
        content = parse_file(path)
        all_content.append(f"【文档 {idx + 1} ({os.path.basename(path)})】\n{content}\n")
    master_content = "\n".join(all_content)

    budget = _resolve_generation_budget(master_content)
    think_enabled = _resolve_think_mode()
    user_msg = _render_user_prompt(master_content, budget, think_enabled)
    logger.info(
        "[Ollama] 生成预算: profile=%s chars=%d num_ctx=%d num_predict=%d think=%s",
        budget.profile_name,
        budget.document_chars,
        budget.context_tokens,
        budget.output_tokens,
        "on" if think_enabled else "off",
    )

    # 初始化 Ollama 客户端
    client_kwargs = {}
    if ollama_host:
        client_kwargs["host"] = ollama_host
    client = ollama.Client(**client_kwargs)

    # 本地模型使用宽松 JSON 模式：不要把完整 JSON Schema 交给解码器硬约束。
    # 27B 量级模型在 schema-constrained decoding 下容易直接空输出或过早中断；
    # 严格校验保留在 _parse_and_validate_situation() 里完成。
    ollama_format = "json"

    last_error: Optional[Exception] = None
    for attempt in range(1, max_retries + 1):
        try:
            logger.info(
                "[Ollama] 提取尝试 %d/%d (model=%s) ...",
                attempt, max_retries, model,
            )
            interrupted = False
            
            interrupted = False

            # ── 流式调用，实时打印输出 ──
            print("\n" + "═" * 60, flush=True, file=sys.stderr)
            print(f"🤖 Ollama [{model}] 正在生成 (流式输出)...", flush=True, file=sys.stderr)
            print("═" * 60, flush=True, file=sys.stderr)

            attempt_user_msg = _build_attempt_user_msg(user_msg, last_error)

            stream = client.chat(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": attempt_user_msg},
                ],
                format=ollama_format,
                stream=True,
                think=think_enabled,
                options={
                    "temperature": 0.2,
                    "top_p": 0.8,
                    "top_k": 20,
                    "repeat_penalty": 1.05,
                    "num_ctx": budget.context_tokens,
                    "num_predict": budget.output_tokens,
                },
            )

            collected_chunks: List[str] = []
            token_count = 0
            t_start = time.time()

            try:
                thinking_started = False
                for chunk in stream:
                    content_text, thinking_text = _extract_ollama_stream_text(chunk)
                    if thinking_text:
                        if not thinking_started:
                            sys.stderr.write("\n[thinking]\n")
                            thinking_started = True
                        token_count += 1
                        sys.stderr.write(thinking_text)
                        sys.stderr.flush()
                    if content_text:
                        if thinking_started:
                            sys.stderr.write("\n[/thinking]\n")
                            thinking_started = False
                        collected_chunks.append(content_text)
                        token_count += 1
                        # 实时打印到控制台
                        sys.stderr.write(content_text)
                        sys.stderr.flush()
            except KeyboardInterrupt:
                print("\n[Ollama] 收到终止信号，提前结束流式生成，将保留已提取的有效数据。", file=sys.stderr, flush=True)
                interrupted = True

            elapsed = time.time() - t_start
            speed = token_count / elapsed if elapsed > 0 else 0

            print(flush=True, file=sys.stderr)
            print("═" * 60, flush=True, file=sys.stderr)
            print(
                f"📊 生成完毕: {token_count} tokens, "
                f"{elapsed:.1f}s, {speed:.1f} tok/s",
                flush=True, file=sys.stderr,
            )
            print("═" * 60 + "\n", flush=True, file=sys.stderr)

            raw_json = "".join(collected_chunks)
            if not raw_json.strip():
                raise ValueError("Ollama 返回空内容")

            # 修复、解析并校验 JSON
            situation = _parse_and_validate_situation(raw_json)
            logger.info(
                "[Ollama] ✅ 成功提取 %d 个目标实体。",
                len(situation.targets),
            )
            return situation

        except Exception as e:
            if interrupted:
                logger.info("[Ollama] 由于用户中断，跳过重试并尝试利用当前已提取的结果。")
                if 'raw_json' in locals() and raw_json.strip():
                    try:
                        return _parse_and_validate_situation(raw_json)
                    except:
                        logger.warning("[Ollama] 中断后的 JSON 修复失败，将返回空目标列表。")
                        return SituationMap(targets=[], enemy_force_type="未知", enemy_force_type_confidence=0, enemy_force_type_basis="用户中止推演")
                return SituationMap(targets=[], enemy_force_type="未知", enemy_force_type_confidence=0, enemy_force_type_basis="用户中止推演")
            last_error = e
            logger.warning("[Ollama] 第 %d 次尝试失败: %s", attempt, str(e))
            if attempt < max_retries:
                logger.info("[Ollama] 将在 2 秒后重试...")
                time.sleep(2)

    raise RuntimeError(
        f"[Ollama] 提取在 {max_retries} 次重试后仍失败: {last_error}"
    )


# ────────────────────────────────────────────────────────
# OpenAI 兼容 API 提取
# ────────────────────────────────────────────────────────

def extract_with_openai_api(
    doc_paths: List[str],
    api_key: str,
    base_url: str = "",
    model: str = "gpt-4o",
    max_retries: int = 3,
) -> SituationMap:
    """使用 OpenAI 兼容 API 从文档中提取结构化态势数据。

    支持 OpenAI、DeepSeek、智谱、通义千问等兼容 OpenAI 接口的服务。
    通过 response_format 约束输出为符合 SituationMap JSON Schema 的结构化数据。

    Args:
        doc_paths: .docx 文件路径列表（第一个为情报文档，第二个为装备手册）。
        api_key: API 密钥。
        base_url: API 基础 URL。
        model: 模型名称。
        max_retries: 最大重试次数。

    Returns:
        SituationMap 实例，包含所有提取的目标实体。

    Raises:
        RuntimeError: 超过最大重试次数后仍解析失败。
    """
    from openai import OpenAI

    if not base_url:
        raise ValueError("使用外部 API 后端时必须提供 base_url。")

    # 解析所有文档
    all_content = []
    import os
    for idx, path in enumerate(doc_paths):
        content = parse_file(path)
        all_content.append(f"【文档 {idx + 1} ({os.path.basename(path)})】\n{content}\n")
    master_content = "\n".join(all_content)

    budget = _resolve_generation_budget(master_content)
    think_enabled = _resolve_think_mode()
    user_msg = _render_user_prompt(master_content, budget, think_enabled)
    logger.info(
        "[API] 生成预算: profile=%s chars=%d max_context_hint=%d max_tokens=%d think=%s",
        budget.profile_name,
        budget.document_chars,
        budget.context_tokens,
        budget.output_tokens,
        "on" if think_enabled else "off",
    )

    client = OpenAI(api_key=api_key, base_url=base_url)

    last_error: Optional[Exception] = None
    for attempt in range(1, max_retries + 1):
        try:
            logger.info(
                "[API] 提取尝试 %d/%d (model=%s, base_url=%s) ...",
                attempt, max_retries, model, base_url,
            )
            
            interrupted = False

            # 确定 response_format (先尝试 json_schema，失败则回退)
            resp_format: dict = {
                "type": "json_schema",
                "json_schema": {
                    "name": "SituationMap",
                    "strict": True,
                    "schema": SituationMap.model_json_schema(),
                },
            }

            # ── 流式调用，实时打印输出 ──
            print("\n" + "═" * 60, flush=True, file=sys.stderr)
            print(
                f"🌐 API [{model}] 正在生成 (流式输出)...",
                flush=True, file=sys.stderr,
            )
            print("═" * 60, flush=True, file=sys.stderr)

            try:
                attempt_user_msg = _build_attempt_user_msg(user_msg, last_error)

                stream = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": attempt_user_msg},
                    ],
                    response_format=resp_format,
                    temperature=0.1,
                    max_tokens=budget.output_tokens,
                    stream=True,
                )
            except Exception:
                # 回退到 json_object 模式
                logger.info(
                    "[API] json_schema 模式不支持，回退到 json_object 模式"
                )
                stream = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": _build_attempt_user_msg(user_msg, last_error)},
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.1,
                    max_tokens=budget.output_tokens,
                    stream=True,
                )

            collected_chunks: List[str] = []
            token_count = 0
            t_start = time.time()

            try:
                reasoning_started = False
                for chunk in stream:
                    delta = chunk.choices[0].delta if chunk.choices else None
                    content_text, reasoning_text = _extract_openai_stream_text(delta)
                    if reasoning_text:
                        if not reasoning_started:
                            sys.stderr.write("\n[reasoning]\n")
                            reasoning_started = True
                        token_count += 1
                        sys.stderr.write(reasoning_text)
                        sys.stderr.flush()
                    if content_text:
                        if reasoning_started:
                            sys.stderr.write("\n[/reasoning]\n")
                            reasoning_started = False
                        collected_chunks.append(content_text)
                        token_count += 1
                        # 实时打印到控制台
                        sys.stderr.write(content_text)
                        sys.stderr.flush()
            except KeyboardInterrupt:
                print("\n[API] 收到终止信号，提前结束流式生成，将保留已提取的有效数据。", file=sys.stderr, flush=True)
                interrupted = True

            elapsed = time.time() - t_start
            speed = token_count / elapsed if elapsed > 0 else 0

            print(flush=True, file=sys.stderr)
            print("═" * 60, flush=True, file=sys.stderr)
            print(
                f"📊 生成完毕: {token_count} tokens, "
                f"{elapsed:.1f}s, {speed:.1f} tok/s",
                flush=True, file=sys.stderr,
            )
            print("═" * 60 + "\n", flush=True, file=sys.stderr)

            raw_json = "".join(collected_chunks)
            if not raw_json.strip():
                raise ValueError("API 返回空内容")

            # 修复、解析并校验 JSON
            situation = _parse_and_validate_situation(raw_json)
            logger.info(
                "[API] ✅ 成功提取 %d 个目标实体。",
                len(situation.targets),
            )
            return situation

        except Exception as e:
            if interrupted:
                logger.info("[API] 由于用户中断，跳过重试并尝试利用当前已提取的结果。")
                if 'raw_json' in locals() and raw_json.strip():
                    try:
                        return _parse_and_validate_situation(raw_json)
                    except:
                        logger.warning("[API] 中断后的 JSON 修复失败，将返回空目标列表。")
                        return SituationMap(targets=[], enemy_force_type="未知", enemy_force_type_confidence=0, enemy_force_type_basis="用户中止推演")
                return SituationMap(targets=[], enemy_force_type="未知", enemy_force_type_confidence=0, enemy_force_type_basis="用户中止推演")
            last_error = e
            logger.warning("[API] 第 %d 次尝试失败: %s", attempt, str(e))
            if attempt < max_retries:
                logger.info("[API] 将在 3 秒后重试...")
                time.sleep(3)

    raise RuntimeError(
        f"[API] 提取在 {max_retries} 次重试后仍失败: {last_error}"
    )


# ────────────────────────────────────────────────────────
# 统一入口
# ────────────────────────────────────────────────────────

def extract_situation(
    doc_paths: List[str],
    backend: LLMBackend,
    model: str,
    api_key: str = "",
    base_url: str = "",
    ollama_host: Optional[str] = None,
    max_retries: int = 3,
) -> SituationMap:
    """统一的态势数据提取入口。

    根据指定的后端类型调度到对应的提取函数。
    本管道强制要求 LLM，无有效配置时将直接报错。

    Args:
        doc_paths: .docx 文件路径列表。
        backend: LLM 后端类型 (ollama / openai_api)。
        model: 模型名称。
        api_key: API 密钥（仅 openai_api 后端需要）。
        base_url: API 基础 URL（仅 openai_api 后端需要）。
        ollama_host: Ollama 服务地址（仅 ollama 后端可选）。
        max_retries: 最大重试次数。

    Returns:
        SituationMap 实例。

    Raises:
        ValueError: 后端配置无效。
        RuntimeError: LLM 提取失败。
        ConnectionError: 无法连接到 LLM 服务。
    """
    if backend == LLMBackend.OLLAMA:
        logger.info("━" * 50)
        logger.info("🤖 使用 Ollama 本地模型: %s", model)
        if ollama_host:
            logger.info("   Ollama 服务地址: %s", ollama_host)
        logger.info("━" * 50)
        return extract_with_ollama(
            doc_paths=doc_paths,
            model=model,
            ollama_host=ollama_host,
            max_retries=max_retries,
        )

    elif backend == LLMBackend.OPENAI_API:
        if not api_key:
            raise ValueError(
                "使用外部 API 后端时必须提供 --api-key 参数。\n"
                "示例: python main.py --backend openai_api "
                "--api-key your-api-key --model gpt-4o ..."
            )
        if not base_url:
            raise ValueError(
                "使用外部 API 后端时必须提供 --base-url 参数或等价环境变量。"
            )
        logger.info("━" * 50)
        logger.info("🌐 使用外部 API: %s (model=%s)", base_url, model)
        logger.info("━" * 50)
        return extract_with_openai_api(
            doc_paths=doc_paths,
            api_key=api_key,
            base_url=base_url,
            model=model,
            max_retries=max_retries,
        )

    else:
        raise ValueError(f"不支持的 LLM 后端类型: {backend}")
